import os
import glob
from docx import Document

def count_terms(file_path, terms):
    counts = {term: 0 for term in terms}
    try:
        doc = Document(file_path)
        for para in doc.paragraphs:
            for term in terms:
                if term in para.text:
                    counts[term] += 1
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for term in terms:
                        if term in cell.text:
                            counts[term] += 1
    except Exception as e:
        print(f"Error reading {file_path}: {e}")
    return counts

if __name__ == "__main__":
    search_pattern = r"practice/data/theory/jungol/중등/2026년_*_중등*2019-2025.docx"
    files = glob.glob(search_pattern)
    files.sort()
    
    terms_to_check = ["유형", "정답", "해설", "문제", "①"]
    
    with open("docx_comprehensive_check.txt", "w", encoding="utf-8") as f:
        for file in files:
            counts = count_terms(file, terms_to_check)
            f.write(f"File: {os.path.basename(file)}\n")
            for term, count in counts.items():
                f.write(f"  '{term}': {count}\n")
            f.write("\n")
