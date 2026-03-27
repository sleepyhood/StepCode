import os
from docx import Document

def extract_text(file_path):
    try:
        doc = Document(file_path)
        with open(f"{os.path.basename(file_path)}.txt", "w", encoding="utf-8") as f:
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    f.write(text + "\n")
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text = cell.text.strip()
                        if text:
                            f.write(text + "\n")
    except Exception as e:
        print(f"Error reading {file_path}: {e}")

if __name__ == "__main__":
    extract_text(r"practice\data\theory\jungol\중등\2026년_1차_중등_수업_교사용_2019-2025.docx")
    extract_text(r"practice\data\theory\jungol\중등\2026년_1차_중등_수업_학생용_2019-2025.docx")
